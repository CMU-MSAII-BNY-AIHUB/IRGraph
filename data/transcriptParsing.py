import docx
import os
import xml.etree.ElementTree as ET
import re
import argparse
import aspose.words as aw

ENDING = ["conclude", "thank you for participating", "no further question",
          "that concludes our call", "that's all the time we have",
          "end of", "thank you for joining",
          "this concludes our presentation", "we are now closing the call",
          "thank you for your questions", "end the call", "thank you all for attending",
          "this ends our session", "no further questions", "closing remarks", "final", 
          "wrap up", "until next time", "goodbye", 
          "appreciate your participation", "thanks to everyone", "signing off", "that's a wrap", 
          "we're done for today", "look forward to our next meeting",  "good day", "closing"]

def remove_empty_columns(arr):
    transposed = list(zip(*arr))

    filtered = [col for col in transposed if any(cell != "" for cell in col)]

    return list(zip(*filtered))

def rtfToDocx(rtf_path, filename):
    doc = aw.Document(os.path.join(rtf_path, filename))
    doc.save(filename.replace(".rtf", ".docx"))
    doc = docx.Document(filename.replace(".rtf", ".docx"))
    return doc

def build_first_table(data):
    data = [list(dict.fromkeys(row)) for row in data]
    data = remove_empty_columns(data)
    root = ET.Element("table",attrib={"id":"0", "name":"Earnings Estimates Comparison Table"})

    time_periods = ET.SubElement(root, "timePeriods")
    for period in data[0][1:]:
        ET.SubElement(time_periods, "period", name=period.replace("-", ""))

    metrics = ET.SubElement(root, "metrics")
    for metric_data in data[2:]:
        metric = ET.SubElement(metrics, "metric", name=metric_data[0])
        for i, value in enumerate(metric_data[1:]):
            period_name = data[0][i + 1]
            value_type = data[1][i + 1]
            ET.SubElement(metric, "value", period=period_name, type=value_type).text = value
    
    return root

def build_second_table(data):
    # clean duplicates   
    data = [list(dict.fromkeys(row)) for row in data]
    data = remove_empty_columns(data)
    
    root = ET.Element("table",attrib={"id":"1","name":"EPS Normalized Comparison Table"})
    

    type_header = data[1]

    time_periods = ET.SubElement(root, "timePeriods") 
    metrics = ET.SubElement(root, "metrics") 
    metric = ET.SubElement(metrics, "metric", name="EPS Normalized")
    for row in data[2:]:
        period = row[0]
        ET.SubElement(time_periods, "period", name=period)
        for i, cell in enumerate(row[1:], 1):
            type_ = type_header[i]  
            if type_ in ["CONSENSUS", "ACTUAL", "SURPRISE"]:
                value_element = ET.SubElement(metric, "value")
                value_element.text = cell
                value_element.set("period", period)
                value_element.set("type", type_)

    return root

def build_third_table(data,company):
    id = 0
    root = ET.Element("CallParticipants")
    speaker_list = {}
         
    current_group = ''
    for row in data[1:]:
        
        row_data = '\n \n \n'.join(row).strip()
        elements = row_data.split('\n \n \n')
        for element in elements:

            lines = element.split('\n')

            if len(lines) == 1 :
                current_group = lines[0].strip()
                
            if len(lines) > 1:

                name = re.sub(r'\s+', ' ', lines[0].strip())
                
                position = lines[1].strip()
                if current_group == "EXECUTIVES":
                    position = company +", " + position
                person_element = ET.SubElement(root, "person", position=position, group=current_group, id = str(id))
                person_element.text = name
                speaker_list[name] = str(id)
                id+=1
    return root, speaker_list

def process_presentation(dialog,speaker_list, name):
    paragraph = dialog.split('\n')

    conversation = ET.Element("section", attrib={"name": name})
    i = 0 
    while i < len(paragraph):
        if re.sub(r'\s+', ' ', paragraph[i].strip())  in speaker_list:
            id = speaker_list[re.sub(r'\s+', ' ', paragraph[i].strip()) ]
            position = paragraph[i+1].strip()
            statement = ET.SubElement(conversation, "statement")
            speaker_element = ET.SubElement(statement, "speaker", id=id, position=position)
            speaker_element.text = re.sub(r'\s+', ' ', paragraph[i].strip()) 
            text = ""
            para = ET.SubElement(speaker_element, "text")
            i += 2
            while i < len(paragraph) and re.sub(r'\s+', ' ', paragraph[i].strip()) not in speaker_list and paragraph[i].strip()!= "Operator":
                if len(paragraph[i].strip()) != 0:
                    text += paragraph[i] + "\n"
                i += 1
            para.text = text.strip()
            
        elif "Operator" in paragraph[i]:
            id = "-1"
            position = ""
            statement = ET.SubElement(conversation, "statement")
            speaker_element = ET.SubElement(statement, "speaker", id=id, position=position)
            speaker_element.text = "Operator"
            text = ""
            para = ET.SubElement(speaker_element, "text")
            i += 1
            while i < len(paragraph) and paragraph[i].strip() not in speaker_list:
                if text != "" or text != None or text != "\n":
                    text += paragraph[i] + "\n"
                i += 1
            para.text = text.strip()
            
        else:
            i += 1
    return conversation


def process_dialog(dialog,speaker_list, name):
    question_id = -1
    followup_id = -1
    end = False
    paragraph = dialog.split('\n')
    cur_question = None
    conversation = ET.Element("section", attrib={"name": name})
    i = 0 
    hasSub = False
    last_question_element = None
    last_question_answered = True
    while i < len(paragraph):
        # print(paragraph[i])
        # print("------------------------")
        if re.sub(r'\s+', ' ', paragraph[i].strip()) in speaker_list:
            id = speaker_list[re.sub(r'\s+', ' ', paragraph[i].strip())]
            position = paragraph[i+1].strip()
            if end:
                context = ET.SubElement(conversation, "ending", id = str(question_id))
                
            elif cur_question == None:
                if last_question_element is not None and not last_question_answered:
                    if last_question_element.tag =="question":
                        question_id-=1
                    last_question_element.tag = "other"
                followup_id = -1
                context = ET.SubElement(conversation, "question", id = str(question_id))
                cur_question = paragraph[i].strip()
                last_question_element = context
                last_question_answered = False
            elif paragraph[i].strip() == cur_question :
                if last_question_element is not None and not last_question_answered:
                    if last_question_element.tag =="question":
                        question_id-=1
                    elif last_question_element.tag =="followQuestion":
                        print(last_question_element.tag)
                        followup_id -=1
                    last_question_element.tag = "other"

                followup_id += 1
                context = ET.SubElement(conversation, "followQuestion", id=str(followup_id),  question_id = str(question_id))
                hasSub = True
                last_question_element = context
                last_question_answered = False
            elif hasSub and paragraph[i].strip()!= cur_question:
                context = ET.SubElement(conversation, "followAnswer", id=str(followup_id),  question_id = str(question_id))
                hasSub = False
                last_question_answered = True
            else:
                context = ET.SubElement(conversation, "answer", id = str(question_id))
                last_question_answered = True
            speaker_element = ET.SubElement(context, "speaker", id=id, position=position)
            speaker_element.text = re.sub(r'\s+', ' ', paragraph[i].strip()) 
            text = ""
            para = ET.SubElement(speaker_element, "text")
            i += 2
            while i < len(paragraph) and re.sub(r'\s+', ' ', paragraph[i].strip()) not in speaker_list and paragraph[i].strip()!= "Operator" and not paragraph[i].startswith("Operator"):
                # print(paragraph[i])
                # print(paragraph[i].startswith("Operator"))
                # print("--------------------------")
                if len(paragraph[i].strip()) != 0:
                    text += paragraph[i] + "\n"
                i += 1
            para.text = text.strip()
            
        elif "Operator" in paragraph[i]:
            if last_question_element is not None and not last_question_answered:
                if last_question_element.tag =="question":
                    question_id-=1
                last_question_element.tag = "other"
            last_question_element = None
            last_question_answered = False
            id = "-1"
            position = ""
            cur_question = None
            hasSub = False
            question_id += 1
            followup_id = -1
            context =ET.SubElement(conversation, "transition") 
            speaker_element = ET.SubElement(context, "speaker", id=id, position=position)
            speaker_element.text = "Operator"
            text = ""
            para = ET.SubElement(speaker_element, "text")
            paragraph[i] = paragraph[i].replace("Operator", "")
            while i < len(paragraph) and paragraph[i].strip() not in speaker_list:
                text += paragraph[i] + "\n"
                i += 1
            para.text = text.strip()
            if "conclude" in para.text:
                context.tag = "ending"
                end = True
                
            
        else:
            i += 1

    return conversation


def prettify(element, indent='    ', level=0):
    
    if element:  
        if not element.text or not element.text.strip():
            element.text = '\n' + indent * (level + 1)
        if not element.tail or not element.tail.strip():
            element.tail = '\n' + indent * level
    else:
        if level and (not element.tail or not element.tail.strip()):
            element.tail = '\n' + indent * level
    
    for subelement in element:
        prettify(subelement, indent, level + 1)


def build_table(doc):
    company = ""
    for i, paragraph in enumerate(doc.paragraphs):

        if i ==2 :
            company = paragraph.text
            break
    tables = []
    for table_index, table in enumerate(doc.tables):
            t = []
            for row_index, row in enumerate(table.rows):
                row_data = []

            
                for cell in row.cells:
                    row_data.append(cell.text.strip())

                if all(element == "" for element in row_data):
                    continue
                t.append(row_data)
            if t== [['']] or t ==[]:
                continue
            tables.append(t)
            
    t1 = build_first_table(tables[0])
    t2 = build_second_table(tables[1])
    sec2,speaker_list = build_third_table(tables[3],company)
    sec1 = ET.Element("section", attrib={"name": "financial tables"})
    sec1.append(t1)
    sec1.append(t2)
    sec2.tag = "section"
    sec2.set("name", "call participants")
    return sec1, sec2, speaker_list

def build_xml(doc):
    sec1, sec2, speaker_list = build_table(doc)
    body = ET.Element("body")
    company = ""
    title = ""
    time = ""
    currency = ""
    note = ""
    QA = None
    presentation = None
    for i, paragraph in enumerate(doc.paragraphs):

        if i ==2 :
            company = paragraph.text
            
        elif i == 3:
            title = paragraph.text
        elif i == 4:
            time = paragraph.text
        elif i == 6:
            currency= paragraph.text
        elif i == 7:
            note= paragraph.text
        
        elif paragraph.text.strip().startswith("Question and Answer"):
            QA = process_dialog(paragraph.text,speaker_list,"Question and Answer")
        elif paragraph.text.strip().startswith("Presentation"):
            
            presentation = process_presentation(paragraph.text,speaker_list,"Presentation ")

  
    header = ET.Element("header")

    ET.SubElement(header, "company").text = company
    ET.SubElement(header, "title").text = title
    ET.SubElement(header, "time").text = time
    ET.SubElement(header, "currency").text = currency
    ET.SubElement(header, "note").text = note
    body.append(sec1)
    body.append(sec2)
    body.append(presentation)
    body.append(QA)
    root = ET.Element("Transcript")
    root.append(header)
    root.append(body)
    return root


def main():
    parser = argparse.ArgumentParser(description='Parse rtf file and convert to XML.')
    parser.add_argument('--file-dir', type=str, required=False, default="transcripts",
                        help='Path to the transcript file to be parsed.')
    parser.add_argument('--save-dir', type=str, required=False, default="xml",
                        help='Path to save the xml')
    args = parser.parse_args()
    file_path = args.file_dir
    save_path = args.save_dir

    print("Listing contents of:", file_path)

    for root, dirs, files in os.walk(os.path.abspath(file_path)):
        for dir_name in dirs:
            print(f"Directory: {dir_name}")
            dir_path = os.path.join(root, dir_name)
            for filename in os.listdir(dir_path):
                file_path = os.path.join(dir_path, filename)
                if os.path.isfile(file_path):
                    print(f"  File: {filename}")
                    doc = rtfToDocx(dir_path,filename)
                    tree_root = build_xml(doc)
                    prettify(tree_root)
                    tree = ET.ElementTree(tree_root)
                    tree.write(os.path.join(save_path,filename.replace(".rtf",".xml")), encoding="utf-8", xml_declaration=True)
                    os.remove(filename.replace(".rtf", ".docx"))
                    
if __name__ == "__main__":
    main()